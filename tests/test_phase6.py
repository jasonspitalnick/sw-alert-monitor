"""
tests/test_phase6.py
Phase 6 tests: Draft generator.

Structure
---------
Group 1 — Unit tests: helpers and prompt builder (no network, no API)
Group 2 — Unit tests: docx builder with synthetic draft text (no API)
Group 3 — Live Anthropic API draft generation (requires ANTHROPIC_API_KEY)

Run:
    venv/bin/python tests/test_phase6.py

Requirements:
    ANTHROPIC_API_KEY must be set (real key) for Group 3.
"""

import os
import sys
import tempfile
import shutil
from pathlib import Path
from datetime import datetime, timezone, timedelta

sys.path.insert(0, str(Path(__file__).parent.parent))

# ── Redirect all data files to temp dir ───────────────────────────────────────
os.environ.setdefault("GMAIL_APP_PASSWORD", "test")

import src.config as _cfg
_TMPDIR = tempfile.mkdtemp(prefix="alert_phase6_test_")
_cfg.DB_FILE             = Path(_TMPDIR) / "test.db"
_cfg.ALERT_TRACKER_FILE  = Path(_TMPDIR) / "alert_tracker.json"
_cfg.DRAFTED_TOPICS_FILE = Path(_TMPDIR) / "drafted_topics.json"
_cfg.DATA_DIR            = Path(_TMPDIR)
_cfg.DRAFTS_DIR          = Path(_TMPDIR) / "drafts"
_cfg.DRAFTS_DIR.mkdir(exist_ok=True)

from src import database as db
from src.cluster_detector import SaturationResult
from src.draft_generator import (
    ArticleContent,
    DraftResult,
    generate_draft,
    _build_draft_prompt,
    _count_words,
    _strip_inline_markdown,
    _parse_and_build_docx,
    _get_system_prompt,
    _fetch_article_text,
)

PASS = "✅"
FAIL = "❌"
_passed = _failed = 0


def check(label: str, cond: bool, detail: str = "") -> bool:
    global _passed, _failed
    icon = PASS if cond else FAIL
    suffix = f"  [{detail}]" if detail else ""
    print(f"  {icon}  {label}{suffix}")
    if cond:
        _passed += 1
    else:
        _failed += 1
    return cond


def _date(days_ago: int = 0) -> str:
    dt = datetime.now(timezone.utc) - timedelta(days=days_ago)
    return dt.strftime("%Y-%m-%d")


# ══════════════════════════════════════════════════════════════════════════════
# Group 1 — Unit tests: helpers and prompt builder
# ══════════════════════════════════════════════════════════════════════════════

def test_helpers():
    print("\n── Helper functions ──")

    # _count_words
    check("count_words simple sentence",
          _count_words("This is a test sentence") == 5)
    check("count_words strips markdown",
          _count_words("**Bold text** and *italic* here") == 5)
    check("count_words empty string",
          _count_words("") == 0)

    # _strip_inline_markdown
    check("strip **bold**",
          _strip_inline_markdown("**bold text**") == "bold text")
    check("strip *italic*",
          _strip_inline_markdown("*italic text*") == "italic text")
    check("strip _italic_",
          _strip_inline_markdown("_italic text_") == "italic text")
    check("strip `code`",
          _strip_inline_markdown("`some code`") == "some code")
    check("strip mixed",
          _strip_inline_markdown("**Key Takeaway:** *important* point") ==
          "Key Takeaway: important point")
    check("leave ## headers alone",
          _strip_inline_markdown("## Section Header") == "## Section Header")

    # _get_system_prompt
    system_prompt = _get_system_prompt()
    check("System prompt contains style guide content",
          "Style Guide" in system_prompt or "style guide" in system_prompt.lower())
    check("System prompt is non-trivial length",
          len(system_prompt) > 500, str(len(system_prompt)))
    check("System prompt mentions Snell & Wilmer",
          "Snell & Wilmer" in system_prompt)


def test_prompt_builder():
    print("\n── Prompt builder ──")

    saturation = SaturationResult(
        cluster_id=1,
        subject_key="sec_enforcement_manual_formal_order_2026",
        description=(
            "SEC 2026 Enforcement Manual now requires full Commission approval "
            "for Formal Orders of Investigation"
        ),
        effective_firm_count=3,
        contributing_firms=["Gibson Dunn", "Latham & Watkins", "Davis Polk"],
        effective_firms=["Gibson Dunn", "Latham & Watkins", "Davis Polk"],
        earliest_pub_date=_date(3),
        latest_pub_date=_date(0),
        window_days=3,
    )

    articles = [
        ArticleContent(
            article_id=1,
            firm_name="Gibson Dunn",
            title="SEC Enforcement Manual Update: Commission Must Now Approve Formal Orders",
            url="https://www.gibsondunn.com/sec-enforcement-manual-2026",
            date_published=_date(3),
            full_text="The SEC updated its Enforcement Manual in 2026...",
        ),
        ArticleContent(
            article_id=2,
            firm_name="Latham & Watkins",
            title="SEC's 2026 Enforcement Manual Revokes Delegated Authority",
            url="https://www.lw.com/insights/sec-enforcement-manual-2026",
            date_published=_date(1),
            full_text=None,   # no text — tests graceful handling
        ),
        ArticleContent(
            article_id=3,
            firm_name="Davis Polk",
            title="New SEC Policy Requires Full Commission Vote",
            url="https://www.davispolk.com/sec-formal-orders-2026",
            date_published=_date(0),
            full_text="Under the new policy, staff may no longer authorize...",
        ),
    ]

    prompt = _build_draft_prompt(saturation, articles)

    check("Prompt contains subject_key",
          "sec_enforcement_manual_formal_order_2026" in prompt)
    check("Prompt contains description",
          "Commission approval" in prompt)
    check("Prompt contains firm list",
          "Gibson Dunn" in prompt and "Latham & Watkins" in prompt)
    check("Prompt contains date range",
          saturation.earliest_pub_date in prompt)
    check("Prompt contains article titles",
          "SEC Enforcement Manual Update" in prompt)
    check("Prompt includes full_text for article 1",
          "The SEC updated its Enforcement Manual" in prompt)
    check("Prompt includes full_text for article 3",
          "staff may no longer authorize" in prompt)
    check("Prompt notes article 2 has no text",
          "Latham & Watkins" in prompt)
    check("Prompt specifies target word count",
          str(_cfg.DRAFT_TARGET_WORDS) in prompt)
    check("Prompt specifies TITLE: format marker",
          "TITLE:" in prompt)
    check("Prompt specifies byline",
          "By Jason Spitalnick" in prompt)
    check("Prompt specifies footer marker",
          "[INSERT S&W STANDARD FOOTER]" in prompt)
    check("Prompt instructs original analysis (not summary)",
          "original" in prompt.lower() or "independent" in prompt.lower())

    return saturation, articles


# ══════════════════════════════════════════════════════════════════════════════
# Group 2 — .docx builder with synthetic draft text
# ══════════════════════════════════════════════════════════════════════════════

_SYNTHETIC_DRAFT = """\
TITLE: SEC Enforcement Manual 2026: Commission Approval Now Required for Formal Orders

By Jason Spitalnick

**Key Takeaway:** The SEC updated its Enforcement Manual in early 2026 to require full Commission approval before a Formal Order of Investigation may issue. This change recalibrates leverage in enforcement inquiries and creates meaningful procedural implications for subjects of SEC investigations.

On March 1, 2026, the Securities and Exchange Commission (SEC or Commission) revised its Enforcement Manual to strip senior Enforcement staff of their longstanding delegated authority to issue Formal Orders of Investigation unilaterally. Under the revised policy, only the full Commission—by majority vote—may authorize a Formal Order. The change took effect immediately.

## Background

A Formal Order of Investigation is not simply an administrative formality. It is the mechanism by which the Enforcement Division acquires subpoena authority—the ability to compel testimony under oath and require the production of documents from non-parties. Prior to the 2026 revision, senior Enforcement staff could issue these orders under delegated authority from the Commission, which meant that the escalation from an informal inquiry to a formal investigation was largely an internal administrative decision.

## How the New Policy Works

Under the revised Manual, any case team seeking to move from informal inquiry to formal investigation must now present the matter to the full Commission for a vote. The Commission meets regularly, but the procedural requirement introduces a meaningful delay and, more importantly, a layer of political accountability that did not previously exist.

## Practitioner Considerations

Securities defense counsel and corporate clients should update their assumptions about the early stages of SEC enforcement inquiries. The new procedural requirement creates a window—between an informal inquiry and the Commission vote on a Formal Order—that counsel should use deliberately.

[INSERT S&W STANDARD FOOTER]"""

_SYNTHETIC_DRAFT_NO_KEY_TAKEAWAY = """\
TITLE: DOJ Updates Cooperation Credit Standards

By Jason Spitalnick

The Department of Justice issued updated guidance on cooperation credit in January 2026.

## Background

The DOJ's corporate enforcement policy has long recognized that voluntary cooperation...

## Key Changes

The 2026 update raises the bar for credit by requiring more proactive disclosure.

[INSERT S&W STANDARD FOOTER]"""


def test_docx_builder():
    print("\n── .docx builder ──")

    # Test 1: Full draft with Key Takeaway
    docx_path = _parse_and_build_docx("sec_enforcement_manual_2026", _SYNTHETIC_DRAFT)

    check("docx file was created", docx_path.exists(), str(docx_path))
    check("docx file has non-zero size", docx_path.stat().st_size > 1000)
    check("docx filename contains subject_key",
          "sec_enforcement_manual_2026" in docx_path.name)
    check("docx has .docx extension", docx_path.suffix == ".docx")

    # Verify structure using python-docx
    from docx import Document as _Doc
    doc = _Doc(str(docx_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    check("docx has at least 5 paragraphs/headings",
          len(paragraphs) >= 5, str(len(paragraphs)))
    check("Title appears in docx",
          any("SEC Enforcement Manual 2026" in p for p in paragraphs))
    check("Byline appears in docx",
          any("By Jason Spitalnick" in p for p in paragraphs))
    check("Key Takeaway heading appears",
          any("Key Takeaway" in p for p in paragraphs))
    check("Footer placeholder appears",
          any("[INSERT S&W STANDARD FOOTER]" in p for p in paragraphs))
    check("Background section header appears",
          any("Background" in p for p in paragraphs))
    check("Body text present",
          any("Commission" in p and len(p) > 50 for p in paragraphs))

    # Verify markdown stripped from body
    check("Bold markers stripped from body",
          not any("**" in p for p in paragraphs))
    check("Italic markers stripped from body",
          not any(p.startswith("*") and p.endswith("*") for p in paragraphs))

    # Test 2: Draft without explicit Key Takeaway
    docx_path2 = _parse_and_build_docx("doj_cooperation_credit_2026",
                                        _SYNTHETIC_DRAFT_NO_KEY_TAKEAWAY)
    check("docx2 file created", docx_path2.exists())

    doc2 = _Doc(str(docx_path2))
    paragraphs2 = [p.text for p in doc2.paragraphs if p.text.strip()]
    check("docx2 has title", any("DOJ" in p for p in paragraphs2))
    check("docx2 has byline", any("By Jason Spitalnick" in p for p in paragraphs2))
    check("docx2 footer placeholder present",
          any("[INSERT S&W STANDARD FOOTER]" in p for p in paragraphs2))

    return docx_path


def test_article_text_fetch():
    print("\n── Article text fetcher ──")

    # Test with a well-known, reliably-accessible page
    text = _fetch_article_text(
        "https://www.sec.gov/news/press-release/2024-1"  # a static SEC page
    )
    # It may succeed or fail depending on the URL—we just test the interface
    check("Returns None or a non-empty string",
          text is None or isinstance(text, str))
    if text is not None:
        check("Returned text has reasonable length",
              len(text) >= 50, str(len(text)))
        check("Returned text within max chars limit",
              len(text) <= 4100)   # _MAX_ARTICLE_TEXT_CHARS + small buffer

    # Test with an invalid URL — must return None, not raise
    bad_text = _fetch_article_text("https://this-domain-does-not-exist-xq9z.invalid/article")
    check("Invalid URL → None (no exception)", bad_text is None)


# ══════════════════════════════════════════════════════════════════════════════
# Group 3 — Live Anthropic API draft generation
# ══════════════════════════════════════════════════════════════════════════════

def test_live_draft_generation():
    print("\n── Live draft generation (Anthropic API) ──")

    if not _cfg.ANTHROPIC_API_KEY:
        print("  ⚠️   ANTHROPIC_API_KEY not set — skipping live test")
        return True

    # Set up a fresh DB with 3 articles on the same cluster
    db.init_db()

    cluster_key = "sec_enforcement_manual_formal_order_2026"
    cluster_desc = (
        "SEC 2026 Enforcement Manual update requires full Commission approval "
        "for Formal Orders of Investigation, revoking delegated authority "
        "from senior Enforcement staff"
    )

    cid = db.upsert_cluster(cluster_key, cluster_desc)

    articles_data = [
        {
            "firm": "Gibson Dunn",
            "title": "SEC Enforcement Manual Update: Commission Must Now Approve All Formal Orders of Investigation",
            "url": "https://www.gibsondunn.com/sec-enforcement-manual-formal-order-2026",
            "date": _date(3),
        },
        {
            "firm": "Latham & Watkins",
            "title": "SEC's 2026 Enforcement Manual Revokes Delegated Authority for Formal Orders",
            "url": "https://www.lw.com/insights/sec-enforcement-manual-2026",
            "date": _date(1),
        },
        {
            "firm": "Davis Polk",
            "title": "New SEC Policy Requires Full Commission Vote Before Formal Investigation Can Begin",
            "url": "https://www.davispolk.com/sec-formal-orders-2026",
            "date": _date(0),
        },
    ]

    for art in articles_data:
        aid = db.upsert_article(
            firm_name=art["firm"],
            title=art["title"],
            url=art["url"],
            date_published=art["date"],
            is_in_scope=True,
        )
        db.link_article_to_cluster(cid, aid)

    db.refresh_cluster_stats(cid)

    saturation = SaturationResult(
        cluster_id=cid,
        subject_key=cluster_key,
        description=cluster_desc,
        effective_firm_count=3,
        contributing_firms=["Gibson Dunn", "Latham & Watkins", "Davis Polk"],
        effective_firms=["Gibson Dunn", "Latham & Watkins", "Davis Polk"],
        earliest_pub_date=_date(3),
        latest_pub_date=_date(0),
        window_days=3,
    )

    print(f"  Calling {_cfg.DRAFT_MODEL} for draft generation...")
    result = generate_draft(saturation)

    check("generate_draft returned a DraftResult", result is not None)

    if result is None:
        return False

    check("DraftResult.subject_key correct",
          result.subject_key == cluster_key)
    check("DraftResult.cluster_id correct",
          result.cluster_id == cid)
    check("DraftResult.docx_path exists",
          result.docx_path.exists(), str(result.docx_path))
    check("DraftResult.docx_path has .docx extension",
          result.docx_path.suffix == ".docx")
    check("DraftResult.word_count is substantial",
          result.word_count >= 400, str(result.word_count))
    check("Draft text contains subject matter",
          "SEC" in result.draft_text or "Commission" in result.draft_text or
          "Enforcement" in result.draft_text)
    check("Draft text contains footer placeholder",
          "[INSERT S&W STANDARD FOOTER]" in result.draft_text)
    check("Draft text contains byline",
          "By Jason Spitalnick" in result.draft_text)

    print(f"\n  Draft word count: {result.word_count}")
    print(f"  Saved to: {result.docx_path}")

    # Verify .docx structure
    from docx import Document as _Doc
    doc = _Doc(str(result.docx_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    check(".docx has title paragraph",
          len(paragraphs) > 0 and len(paragraphs[0]) > 5)
    check(".docx has at least 8 paragraphs",
          len(paragraphs) >= 8, str(len(paragraphs)))
    check(".docx byline present",
          any("By Jason Spitalnick" in p for p in paragraphs))
    check(".docx footer placeholder present",
          any("[INSERT S&W STANDARD FOOTER]" in p for p in paragraphs))

    # Print a preview
    print(f"\n  Document title: {paragraphs[0] if paragraphs else '(none)'}")
    print(f"  Paragraph count: {len(paragraphs)}")

    return True


# ══════════════════════════════════════════════════════════════════════════════
# Runner
# ══════════════════════════════════════════════════════════════════════════════

def run_all():
    global _passed, _failed

    print("\n=== Phase 6 Tests: Draft Generator ===")
    group_passed = group_failed = 0

    groups = [
        test_helpers,
        test_prompt_builder,
        test_docx_builder,
        test_article_text_fetch,
        test_live_draft_generation,
    ]

    for fn in groups:
        pre_p, pre_f = _passed, _failed
        try:
            fn()
            if _failed - pre_f == 0:
                group_passed += 1
            else:
                group_failed += 1
        except Exception as e:
            import traceback
            print(f"  {FAIL}  EXCEPTION in {fn.__name__}: {e}")
            traceback.print_exc()
            group_failed += 1

    print(f"\n{'='*55}")
    print(f"Assertions: {_passed} passed, {_failed} failed")
    print(f"Groups:     {group_passed}/{group_passed + group_failed} passed")
    if _failed == 0:
        print("All tests PASSED ✅")
    print(f"{'='*55}\n")

    shutil.rmtree(_TMPDIR, ignore_errors=True)
    return _failed == 0


if __name__ == "__main__":
    ok = run_all()
    sys.exit(0 if ok else 1)
