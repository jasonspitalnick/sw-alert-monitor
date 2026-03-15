"""
tests/test_phase7.py
Phase 7 tests: Email delivery.

Structure
---------
Group 1 — Unit tests: subject line, body builders, MIME assembly (no SMTP)
Group 2 — Live SMTP credential check (connect + auth only, no email sent)
           Requires GMAIL_APP_PASSWORD to be set.
Group 3 — Live full send (requires GMAIL_APP_PASSWORD + SEND_TEST_EMAIL=1)
           Set env var SEND_TEST_EMAIL=1 to enable.

Run:
    venv/bin/python tests/test_phase7.py
    SEND_TEST_EMAIL=1 venv/bin/python tests/test_phase7.py   # also sends email

Requirements:
    GMAIL_APP_PASSWORD must be set for Groups 2 and 3.
    SEND_TEST_EMAIL=1 required to actually deliver an email (Group 3).
"""

import os
import sys
import tempfile
import shutil
from pathlib import Path
from datetime import datetime, timezone, timedelta
from email.message import EmailMessage

sys.path.insert(0, str(Path(__file__).parent.parent))

# ── Redirect data files to temp dir ───────────────────────────────────────────
import src.config as _cfg
_TMPDIR = tempfile.mkdtemp(prefix="alert_phase7_test_")
_cfg.DB_FILE             = Path(_TMPDIR) / "test.db"
_cfg.ALERT_TRACKER_FILE  = Path(_TMPDIR) / "alert_tracker.json"
_cfg.DRAFTED_TOPICS_FILE = Path(_TMPDIR) / "drafted_topics.json"
_cfg.DATA_DIR            = Path(_TMPDIR)
_cfg.DRAFTS_DIR          = Path(_TMPDIR) / "drafts"
_cfg.DRAFTS_DIR.mkdir(exist_ok=True)

from src import database as db
from src.cluster_detector import SaturationResult
from src.draft_generator import DraftResult
from src.email_delivery import (
    AlertEmail,
    build_subject,
    build_body_text,
    build_body_html,
    build_alert_email,
    _build_mime_message,
    check_smtp_credentials,
    send_alert_email,
)

PASS = "✅"
FAIL = "❌"
_passed = _failed = 0


def check(label: str, cond: bool, detail: str = "") -> bool:
    global _passed, _failed
    icon = PASS if cond else FAIL
    suffix = f"  [{detail}]" if detail else ""
    print(f"  {icon}  {label}{suffix}")
    if cond:
        _passed += 1
    else:
        _failed += 1
    return cond


def _date(days_ago: int = 0) -> str:
    dt = datetime.now(timezone.utc) - timedelta(days=days_ago)
    return dt.strftime("%Y-%m-%d")


# ── Fixtures ───────────────────────────────────────────────────────────────────

def _make_saturation() -> SaturationResult:
    return SaturationResult(
        cluster_id=99,
        subject_key="sec_enforcement_manual_formal_order_2026",
        description=(
            "SEC 2026 Enforcement Manual now requires full Commission approval "
            "for Formal Orders of Investigation"
        ),
        effective_firm_count=3,
        contributing_firms=["Gibson Dunn", "Latham & Watkins", "Davis Polk"],
        effective_firms=["Gibson Dunn", "Latham & Watkins", "Davis Polk"],
        earliest_pub_date=_date(3),
        latest_pub_date=_date(0),
        window_days=3,
    )


def _make_docx(filename: str = "test_draft.docx") -> Path:
    """Create a minimal valid .docx file for attachment testing."""
    from docx import Document
    path = _cfg.DRAFTS_DIR / filename
    doc = Document()
    doc.add_heading("Test Draft Alert", level=0)
    doc.add_paragraph("By Jason Spitalnick")
    doc.add_paragraph("This is a test draft.")
    doc.add_paragraph("[INSERT S&W STANDARD FOOTER]")
    doc.save(str(path))
    return path


def _make_draft_result(docx_path: Path) -> DraftResult:
    return DraftResult(
        cluster_id=99,
        subject_key="sec_enforcement_manual_formal_order_2026",
        docx_path=docx_path,
        draft_text="Test draft content by Jason Spitalnick.",
        word_count=850,
    )


# ══════════════════════════════════════════════════════════════════════════════
# Group 1 — Unit tests: subject, body, MIME assembly
# ══════════════════════════════════════════════════════════════════════════════

def test_subject_line():
    print("\n── Subject line builder ──")

    sat = _make_saturation()
    subject = build_subject(sat)

    check("Subject starts with [DRAFT ALERT]",
          subject.startswith("[DRAFT ALERT]"))
    check("Subject contains cluster description",
          "Commission approval" in subject or "Formal Orders" in subject)
    check("Subject contains a date",
          any(m in subject for m in [
              "January", "February", "March", "April", "May", "June",
              "July", "August", "September", "October", "November", "December",
          ]))
    check("Subject contains em-dash separator",
          " — " in subject)
    check("Subject length is reasonable",
          20 < len(subject) < 250, str(len(subject)))

    # Long description should be truncated
    long_sat = SaturationResult(
        cluster_id=1,
        subject_key="test",
        description="A" * 200,
        effective_firm_count=3,
        contributing_firms=[],
        effective_firms=[],
        earliest_pub_date=_date(2),
        latest_pub_date=_date(0),
        window_days=2,
    )
    long_subject = build_subject(long_sat)
    check("Long description truncated in subject",
          len(long_subject) < 300, str(len(long_subject)))


def test_body_builders():
    print("\n── Body builders ──")

    # Set up DB articles for the test cluster
    db.init_db()
    sat = _make_saturation()
    cid = db.upsert_cluster(sat.subject_key, sat.description)
    # Patch cluster_id so it matches the DB row
    sat = SaturationResult(
        cluster_id=cid,
        subject_key=sat.subject_key,
        description=sat.description,
        effective_firm_count=sat.effective_firm_count,
        contributing_firms=sat.contributing_firms,
        effective_firms=sat.effective_firms,
        earliest_pub_date=sat.earliest_pub_date,
        latest_pub_date=sat.latest_pub_date,
        window_days=sat.window_days,
    )

    art_ids = []
    for i, firm in enumerate(["Gibson Dunn", "Latham & Watkins", "Davis Polk"]):
        aid = db.upsert_article(
            firm_name=firm,
            title=f"SEC Enforcement Manual Update — {firm}",
            url=f"https://example.com/{firm.lower().replace(' ', '-')}-sec",
            date_published=_date(i),
            is_in_scope=True,
        )
        db.link_article_to_cluster(cid, aid)
        art_ids.append(aid)

    docx_path = _make_docx()
    draft = _make_draft_result(docx_path)
    articles = db.get_cluster_articles(cid)

    # Plain text body
    text = build_body_text(sat, draft, articles)
    check("Plain text contains triggering subject",
          "Commission approval" in text or "Formal Orders" in text)
    check("Plain text contains firm names",
          "Gibson Dunn" in text and "Latham & Watkins" in text)
    check("Plain text contains article links",
          "https://example.com" in text)
    check("Plain text contains date range",
          sat.earliest_pub_date in text)
    check("Plain text confirms gap status",
          "Snell & Wilmer" in text and ("not published" in text or "gap" in text.lower()))
    check("Plain text mentions draft attachment",
          "first draft" in text.lower() or "draft" in text.lower())
    check("Plain text mentions docx filename",
          draft.docx_path.name in text)
    check("Plain text mentions word count",
          str(draft.word_count) in text)

    # HTML body
    html = build_body_html(sat, draft, articles)
    check("HTML is well-formed (has <html> tags)",
          "<html" in html.lower() and "</html>" in html.lower())
    check("HTML contains firm names",
          "Gibson Dunn" in html and "Davis Polk" in html)
    check("HTML contains article links as <a> tags",
          '<a href="https://example.com' in html)
    check("HTML contains gap status section",
          "Snell &amp; Wilmer" in html or "Snell & Wilmer" in html)
    check("HTML contains draft warning",
          "first draft" in html.lower())
    check("HTML mentions docx filename",
          draft.docx_path.name in html)

    return sat, draft, cid


def test_email_assembly():
    print("\n── Email assembly (build_alert_email) ──")

    db.init_db()
    sat = _make_saturation()
    cid = db.upsert_cluster(sat.subject_key, sat.description)
    sat = SaturationResult(
        cluster_id=cid, subject_key=sat.subject_key,
        description=sat.description,
        effective_firm_count=sat.effective_firm_count,
        contributing_firms=sat.contributing_firms,
        effective_firms=sat.effective_firms,
        earliest_pub_date=sat.earliest_pub_date,
        latest_pub_date=sat.latest_pub_date,
        window_days=sat.window_days,
    )
    for i, firm in enumerate(["Gibson Dunn", "Latham & Watkins", "Davis Polk"]):
        aid = db.upsert_article(
            firm_name=firm,
            title=f"SEC Update — {firm}",
            url=f"https://example.com/{i}",
            date_published=_date(i),
            is_in_scope=True,
        )
        db.link_article_to_cluster(cid, aid)

    docx_path = _make_docx("assembly_test.docx")
    draft = _make_draft_result(docx_path)

    alert = build_alert_email(sat, draft)

    check("AlertEmail.to_addr set",
          "@" in alert.to_addr, alert.to_addr)
    check("AlertEmail.from_addr set",
          "@" in alert.from_addr, alert.from_addr)
    check("AlertEmail.subject correct format",
          alert.subject.startswith("[DRAFT ALERT]"))
    check("AlertEmail.body_text non-empty",
          len(alert.body_text) > 100, str(len(alert.body_text)))
    check("AlertEmail.body_html non-empty",
          len(alert.body_html) > 200, str(len(alert.body_html)))
    check("AlertEmail.attachment_path correct",
          alert.attachment_path == docx_path)
    check("AlertEmail.attachment_name correct",
          alert.attachment_name == docx_path.name)


def test_mime_message():
    print("\n── MIME message builder ──")

    docx_path = _make_docx("mime_test.docx")
    alert = AlertEmail(
        to_addr="recipient@example.com",
        from_addr="sender@example.com",
        subject="[DRAFT ALERT] Test Subject — March 14, 2026",
        body_text="Plain text body content here.",
        body_html="<html><body><p>HTML body content here.</p></body></html>",
        attachment_path=docx_path,
        attachment_name=docx_path.name,
    )

    msg = _build_mime_message(alert)

    check("Message is an EmailMessage",
          isinstance(msg, EmailMessage))
    check("To header set correctly",
          msg["To"] == "recipient@example.com")
    check("From header set correctly",
          msg["From"] == "sender@example.com")
    check("Subject header set correctly",
          msg["Subject"] == "[DRAFT ALERT] Test Subject — March 14, 2026")

    # Serialize the message and check for expected content
    raw = msg.as_string()
    check("Raw message contains plain text body",
          "Plain text body content here" in raw)
    check("Raw message contains HTML body",
          "HTML body content here" in raw)
    check("Raw message contains attachment filename",
          docx_path.name in raw)
    check("Raw message is multipart",
          "multipart" in raw.lower() or msg.is_multipart())
    check("Message has positive size",
          len(raw) > 500, str(len(raw)))

    # Test with missing attachment (should not raise)
    missing_path = _cfg.DRAFTS_DIR / "nonexistent.docx"
    alert_no_attach = AlertEmail(
        to_addr="r@e.com", from_addr="s@e.com",
        subject="Test", body_text="text", body_html="<p>html</p>",
        attachment_path=missing_path,
        attachment_name="nonexistent.docx",
    )
    msg2 = _build_mime_message(alert_no_attach)
    check("Missing attachment handled gracefully (no exception)",
          isinstance(msg2, EmailMessage))


# ══════════════════════════════════════════════════════════════════════════════
# Group 2 — Live SMTP credential check (connect + auth, no send)
# ══════════════════════════════════════════════════════════════════════════════

def test_smtp_credentials():
    print("\n── SMTP credential check (no email sent) ──")

    if not _cfg.GMAIL_APP_PASSWORD:
        print("  ⚠️   GMAIL_APP_PASSWORD not set — skipping SMTP test")
        return True

    result = check_smtp_credentials()
    check("SMTP credentials valid (connect + auth succeeded)", result)

    if not result:
        print("  ℹ️   Verify GMAIL_APP_PASSWORD is a valid Gmail app password")
        print("       (Not your regular Gmail password — must be an app password)")

    return result


# ══════════════════════════════════════════════════════════════════════════════
# Group 3 — Live send (SEND_TEST_EMAIL=1 required)
# ══════════════════════════════════════════════════════════════════════════════

def test_live_send():
    print("\n── Live email send (SEND_TEST_EMAIL=1) ──")

    if not os.environ.get("SEND_TEST_EMAIL"):
        print("  ⚠️   SEND_TEST_EMAIL not set — skipping live send test")
        print("       (Set SEND_TEST_EMAIL=1 to actually send a test email)")
        return True

    if not _cfg.GMAIL_APP_PASSWORD:
        print("  ⚠️   GMAIL_APP_PASSWORD not set — skipping live send test")
        return True

    db.init_db()
    sat = _make_saturation()
    cid = db.upsert_cluster(sat.subject_key, sat.description)
    sat = SaturationResult(
        cluster_id=cid, subject_key=sat.subject_key,
        description=sat.description,
        effective_firm_count=sat.effective_firm_count,
        contributing_firms=sat.contributing_firms,
        effective_firms=sat.effective_firms,
        earliest_pub_date=sat.earliest_pub_date,
        latest_pub_date=sat.latest_pub_date,
        window_days=sat.window_days,
    )
    for i, firm in enumerate(["Gibson Dunn", "Latham & Watkins", "Davis Polk"]):
        aid = db.upsert_article(
            firm_name=firm,
            title=f"SEC Enforcement Manual Update — {firm}",
            url=f"https://example.com/{firm.lower().replace(' ', '-')}",
            date_published=_date(i),
            is_in_scope=True,
        )
        db.link_article_to_cluster(cid, aid)

    docx_path = _make_docx("live_send_test.docx")
    draft = _make_draft_result(docx_path)

    # Modify subject to make it clearly a test
    alert = build_alert_email(sat, draft)
    alert = AlertEmail(
        to_addr=alert.to_addr,
        from_addr=alert.from_addr,
        subject=f"[TEST] {alert.subject}",
        body_text=f"⚠️ THIS IS A TEST EMAIL\n\n{alert.body_text}",
        body_html=alert.body_html.replace(
            "Alert Monitor —",
            "Alert Monitor [TEST MODE] —",
        ),
        attachment_path=alert.attachment_path,
        attachment_name=alert.attachment_name,
    )

    print(f"  Sending test email to {alert.to_addr}...")
    success = send_alert_email(alert)

    check("Email sent successfully", success)

    if success:
        print(f"  ✉️   Test email delivered to {alert.to_addr}")
        print(f"       Subject: {alert.subject}")

    return success


# ══════════════════════════════════════════════════════════════════════════════
# Runner
# ══════════════════════════════════════════════════════════════════════════════

def run_all():
    global _passed, _failed

    print("\n=== Phase 7 Tests: Email Delivery ===")
    group_passed = group_failed = 0

    groups = [
        test_subject_line,
        test_body_builders,
        test_email_assembly,
        test_mime_message,
        test_smtp_credentials,
        test_live_send,
    ]

    for fn in groups:
        pre_p, pre_f = _passed, _failed
        try:
            fn()
            if _failed - pre_f == 0:
                group_passed += 1
            else:
                group_failed += 1
        except Exception as e:
            import traceback
            print(f"  {FAIL}  EXCEPTION in {fn.__name__}: {e}")
            traceback.print_exc()
            group_failed += 1

    print(f"\n{'='*55}")
    print(f"Assertions: {_passed} passed, {_failed} failed")
    print(f"Groups:     {group_passed}/{group_passed + group_failed} passed")
    if _failed == 0:
        print("All tests PASSED ✅")
    print(f"{'='*55}\n")

    shutil.rmtree(_TMPDIR, ignore_errors=True)
    return _failed == 0


if __name__ == "__main__":
    ok = run_all()
    sys.exit(0 if ok else 1)
