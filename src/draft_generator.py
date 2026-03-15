"""
src/draft_generator.py
Phase 6 — AI-powered draft generation and .docx export.

Flow
----
1. Receive a SaturationResult (saturated cluster) from the caller.
2. Fetch the linked articles from the DB.
3. Optionally enrich each article by fetching its full text (httpx + BS4).
4. Build a user-message prompt with the triggering subject and competitor
   article content.
5. Call claude-opus-4-6 with STYLE_GUIDE.md as the system prompt.
6. Parse the response and write a formatted .docx to DRAFTS_DIR.
7. Return a DraftResult with the path, raw text, and word count.

Prompt design (per spec)
------------------------
* System prompt  = full STYLE_GUIDE.md content (± a one-line framing header).
  The style guide governs voice, structure, anti-LLM principles, etc.
* User message   = triggering subject + competitor article content +
  explicit format instructions.  Claude is told to write an *original*
  analysis of the underlying development, not a summary of competitors.
"""

import logging
import re
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

import anthropic
import httpx
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tenacity import (
    retry,
    retry_if_exception_type,
    stop_after_attempt,
    wait_exponential,
)

from src.config import (
    ANTHROPIC_API_KEY,
    DRAFT_MODEL,
    DRAFT_TARGET_WORDS,
    DRAFT_BYLINE,
    STYLE_GUIDE_FILE,
    DRAFTS_DIR,
)
from src import database as db
from src.cluster_detector import SaturationResult

logger = logging.getLogger(__name__)

# ── Constants ──────────────────────────────────────────────────────────────────

_USER_AGENT = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36"
)

# Max characters of article body text to include per competitor article.
# Large enough to give Claude real content; small enough to keep costs down.
_MAX_ARTICLE_TEXT_CHARS: int = 4_000

# HTTP fetch timeout for article text retrieval
_FETCH_TIMEOUT: float = 10.0


# ── Data structures ────────────────────────────────────────────────────────────

@dataclass
class ArticleContent:
    """A competitor article with its fetched text (or just metadata)."""
    article_id: int
    firm_name: str
    title: str
    url: str
    date_published: Optional[str]
    full_text: Optional[str] = None   # fetched body text, may be None


@dataclass
class DraftResult:
    cluster_id: int
    subject_key: str
    docx_path: Path
    draft_text: str    # raw text returned by Claude
    word_count: int


# ── Article text fetching ──────────────────────────────────────────────────────

def _fetch_article_text(url: str) -> Optional[str]:
    """
    Fetch and extract the main body text of a competitor article.

    Tries candidate selectors in priority order; falls back to the full
    body text.  Returns at most _MAX_ARTICLE_TEXT_CHARS characters, or
    None on any error.
    """
    try:
        resp = httpx.get(
            url,
            headers={"User-Agent": _USER_AGENT},
            follow_redirects=True,
            timeout=_FETCH_TIMEOUT,
        )
        resp.raise_for_status()
    except Exception as exc:
        logger.debug("Could not fetch article text from %s: %s", url, exc)
        return None

    try:
        soup = BeautifulSoup(resp.text, "lxml")

        # Remove boilerplate containers
        for tag in soup.find_all(
            ["nav", "header", "footer", "aside", "script", "style",
             "noscript", "figure"]
        ):
            tag.decompose()

        # Try candidates in order — first non-trivial match wins
        candidates = [
            soup.find("article"),
            soup.find("main"),
            soup.find(class_=re.compile(
                r"article[-_]?(?:body|content|text)|"
                r"post[-_]?content|entry[-_]?content|"
                r"rich[-_]?text|body[-_]?text|content[-_]?area|"
                r"page[-_]?content|publication[-_]?content",
                re.IGNORECASE,
            )),
            soup.find("body"),
        ]

        for candidate in candidates:
            if candidate is None:
                continue
            text = candidate.get_text(separator="\n", strip=True)
            # Require at least 200 chars to be a real article body
            if len(text) >= 200:
                return text[:_MAX_ARTICLE_TEXT_CHARS]

    except Exception as exc:
        logger.debug("Error parsing article text from %s: %s", url, exc)

    return None


def _enrich_cluster_articles(cluster_id: int) -> list[ArticleContent]:
    """
    Fetch articles linked to the cluster from the DB, then attempt to
    retrieve the full text of each via HTTP.  Returns a list of
    ArticleContent objects sorted by date_published ascending.
    """
    rows = db.get_cluster_articles(cluster_id)
    enriched: list[ArticleContent] = []

    for row in rows:
        # Use stored full_text if available, otherwise fetch from URL
        text: Optional[str] = row["full_text"]
        if not text and row["url"]:
            logger.debug("Fetching article text: %s", row["url"])
            text = _fetch_article_text(row["url"])

        enriched.append(ArticleContent(
            article_id=row["id"],
            firm_name=row["firm_name"],
            title=row["title"],
            url=row["url"],
            date_published=row["date_published"],
            full_text=text,
        ))

    return sorted(
        enriched,
        key=lambda a: a.date_published or "0000-00-00",
    )


# ── Prompt construction ────────────────────────────────────────────────────────

def _get_system_prompt() -> str:
    """
    Read STYLE_GUIDE.md and return it as the system prompt.
    A brief framing header is prepended so Claude understands its role.
    """
    try:
        style_guide = STYLE_GUIDE_FILE.read_text(encoding="utf-8")
    except Exception as exc:
        logger.error("Could not read STYLE_GUIDE.md: %s", exc)
        style_guide = "(Style guide unavailable — apply standard legal writing conventions.)"

    return (
        "You are a legal writing assistant for Snell & Wilmer's white collar "
        "and securities practice group.  Draft client alerts exactly as "
        "prescribed by the style guide below.\n\n"
        + style_guide
    )


def _build_draft_prompt(
    saturation: SaturationResult,
    articles: list[ArticleContent],
) -> str:
    """
    Build the user-message prompt passed to the draft model.

    The prompt supplies the triggering subject, the competitor coverage as
    context, and explicit formatting instructions.  Claude is directed to
    produce an *original* independent analysis — not a synthesis of the
    competitor articles.
    """
    firm_list = ", ".join(saturation.effective_firms)
    date_range = (
        f"{saturation.earliest_pub_date} to {saturation.latest_pub_date} "
        f"({saturation.window_days} days)"
    )

    # Build the competitor articles block
    article_blocks: list[str] = []
    for i, art in enumerate(articles, 1):
        block_lines = [
            f"--- Competitor Article {i} ---",
            f"Firm:  {art.firm_name}",
            f"Title: {art.title}",
            f"Date:  {art.date_published or 'unknown'}",
            f"URL:   {art.url}",
        ]
        if art.full_text:
            block_lines.append(f"\nExcerpt:\n{art.full_text}")
        block_lines.append("")
        article_blocks.append("\n".join(block_lines))

    competitor_section = "\n".join(article_blocks)

    return f"""\
TRIGGERING DEVELOPMENT:
  Subject Key:   {saturation.subject_key}
  Description:   {saturation.description}
  Date Range:    {date_range}
  Firms Covered: {firm_list}

COMPETITOR COVERAGE ({len(articles)} article{"s" if len(articles) != 1 else ""}):
Use the following competitor alerts to understand the scope of the legal \
development and the key issues it raises.  Do not synthesize, summarize, or \
paraphrase these articles.  Write an independent, original analysis of the \
underlying development.

{competitor_section}
DRAFTING INSTRUCTIONS:
1. Draft an original client alert of approximately {DRAFT_TARGET_WORDS} words.
2. This is NOT a summary of the competitor articles.  It is an independent \
analysis of the underlying legal development, written for Snell & Wilmer \
clients who need to understand the development and its practitioner implications.
3. Use the competitor coverage to understand the factual record and key issues, \
then produce your own analysis.
4. Follow all instructions in the style guide (your system prompt) precisely.

FORMAT YOUR RESPONSE EXACTLY AS FOLLOWS (use these exact markers):

TITLE: [Your article title]

{DRAFT_BYLINE}

[Optional: if the development is amenable to a clear 2-4 sentence distillation:]
**Key Takeaway:** [2-4 sentences]

[Body of the article — approximately {DRAFT_TARGET_WORDS} words, with \
## Section Header for any section breaks.  Do not use # headings.]

[INSERT S&W STANDARD FOOTER]"""


# ── Anthropic API call ─────────────────────────────────────────────────────────

_draft_client: Optional[anthropic.Anthropic] = None


def _get_client() -> anthropic.Anthropic:
    global _draft_client
    if _draft_client is None:
        import src.config as _live_cfg
        key = _live_cfg.ANTHROPIC_API_KEY or ANTHROPIC_API_KEY
        _draft_client = anthropic.Anthropic(api_key=key)
    return _draft_client


@retry(
    retry=retry_if_exception_type(
        (anthropic.APIConnectionError, anthropic.RateLimitError)
    ),
    wait=wait_exponential(multiplier=2, min=5, max=60),
    stop=stop_after_attempt(3),
)
def _call_draft_api(system_prompt: str, user_prompt: str) -> str:
    """
    Call claude-opus-4-6 for draft generation.
    Returns the raw text response.
    """
    client = _get_client()
    response = client.messages.create(
        model=DRAFT_MODEL,
        max_tokens=4096,
        system=system_prompt,
        messages=[{"role": "user", "content": user_prompt}],
    )
    return response.content[0].text.strip()


# ── Text utilities ─────────────────────────────────────────────────────────────

def _count_words(text: str) -> int:
    """Count words in plain text (strips markdown markers first)."""
    clean = re.sub(r"[*_`#\[\]()]+", " ", text)
    return len(clean.split())


def _strip_inline_markdown(text: str) -> str:
    """
    Remove common inline markdown markers (**bold**, *italic*, _italic_,
    `code`) from a string, returning plain text.
    Does NOT remove ## headers — those are handled structurally.
    """
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)   # **bold**
    text = re.sub(r"\*(.+?)\*",     r"\1", text)   # *italic*
    text = re.sub(r"_(.+?)_",       r"\1", text)   # _italic_
    text = re.sub(r"`(.+?)`",       r"\1", text)   # `code`
    return text


# ── .docx builder ──────────────────────────────────────────────────────────────

def _parse_and_build_docx(subject_key: str, draft_text: str) -> Path:
    """
    Parse the formatted draft text from Claude and write a .docx file.

    Expected format markers:
      TITLE: <text>           → Heading 1 / Title style
      By Jason Spitalnick     → Centered italic byline
      **Key Takeaway:** <...> → "Key Takeaway" Heading 2 + body paragraph
      ## <text>               → Heading 2 section header
      [INSERT S&W STANDARD FOOTER] → Italic placeholder at end
      All other text          → Normal body paragraph

    Saves to DRAFTS_DIR/<subject_key>_<YYYYMMDD_HHMMSS>.docx.
    Returns the saved path.
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # ── Parse draft text ──────────────────────────────────────────────────────
    lines = draft_text.splitlines()
    title_written = False
    byline_written = False
    para_buffer: list[str] = []

    def _flush_buffer():
        """Write any accumulated paragraph lines as a Normal paragraph."""
        text = " ".join(para_buffer).strip()
        if text:
            # Key Takeaway inline detection: paragraph starts with **Key Takeaway:**
            kt_match = re.match(
                r"\*\*Key Takeaway:\*\*\s*(.*)",
                text,
                re.IGNORECASE,
            )
            if kt_match:
                doc.add_heading("Key Takeaway", level=2)
                kt_text = _strip_inline_markdown(kt_match.group(1)).strip()
                if kt_text:
                    p = doc.add_paragraph(kt_text)
                    p.runs[0].italic = True
            elif text == "[INSERT S&W STANDARD FOOTER]":
                p = doc.add_paragraph(text)
                p.runs[0].italic = True
            else:
                doc.add_paragraph(_strip_inline_markdown(text))
        para_buffer.clear()

    for raw_line in lines:
        line = raw_line.rstrip()

        # ── TITLE: marker (first occurrence) ──────────────────────────────────
        if not title_written and line.upper().startswith("TITLE:"):
            title_text = line[6:].strip()
            if not title_text and para_buffer:
                # Title text might be on the next line — handled below
                title_text = " ".join(para_buffer).strip()
                para_buffer.clear()
            if title_text:
                doc.add_heading(title_text, level=0)   # "Title" style
                title_written = True
            continue

        # ── If we haven't found TITLE: yet, the very first non-empty line IS
        #    the title (Claude occasionally omits the marker)
        if not title_written and line.strip():
            _flush_buffer()
            doc.add_heading(_strip_inline_markdown(line.strip()), level=0)
            title_written = True
            continue

        # ── Byline ────────────────────────────────────────────────────────────
        if not byline_written and DRAFT_BYLINE.lower() in line.lower():
            _flush_buffer()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(DRAFT_BYLINE)
            run.italic = True
            byline_written = True
            continue

        # ── Section header  ## Title ──────────────────────────────────────────
        if line.startswith("## "):
            _flush_buffer()
            doc.add_heading(
                _strip_inline_markdown(line[3:].strip()),
                level=2,
            )
            continue

        # ── Also handle # Title (level 1 — treat as level 2 sub-heading) ──────
        if line.startswith("# ") and not line.upper().startswith("# TITLE"):
            _flush_buffer()
            doc.add_heading(
                _strip_inline_markdown(line[2:].strip()),
                level=2,
            )
            continue

        # ── Footer placeholder ─────────────────────────────────────────────────
        if "[INSERT S&W STANDARD FOOTER]" in line:
            _flush_buffer()
            p = doc.add_paragraph("[INSERT S&W STANDARD FOOTER]")
            if p.runs:
                p.runs[0].italic = True
            continue

        # ── Empty line → flush paragraph ──────────────────────────────────────
        if not line.strip():
            _flush_buffer()
            continue

        # ── Regular text → accumulate ─────────────────────────────────────────
        para_buffer.append(line.strip())

    # Flush anything remaining
    _flush_buffer()

    # ── Ensure footer placeholder appears even if Claude omitted it ───────────
    # Check if footer was written; if not, add it
    footer_present = any(
        "[INSERT S&W STANDARD FOOTER]" in (p.text or "")
        for p in doc.paragraphs
    )
    if not footer_present:
        p = doc.add_paragraph("[INSERT S&W STANDARD FOOTER]")
        if p.runs:
            p.runs[0].italic = True

    # ── Save .docx ────────────────────────────────────────────────────────────
    DRAFTS_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    safe_key = re.sub(r"[^a-z0-9_]", "_", subject_key.lower())[:60]
    filename = f"{safe_key}_{timestamp}.docx"
    docx_path = DRAFTS_DIR / filename

    doc.save(str(docx_path))
    logger.info("Saved draft to: %s", docx_path)
    return docx_path


# ── Public entry point ─────────────────────────────────────────────────────────

def generate_draft(saturation: SaturationResult) -> Optional[DraftResult]:
    """
    Generate a client alert draft for a saturated cluster.

    Parameters
    ----------
    saturation : SaturationResult
        The cluster that has reached the saturation threshold.

    Returns
    -------
    DraftResult | None
        Contains the .docx path, raw draft text, and word count.
        Returns None if the API call fails unrecoverably.
    """
    logger.info(
        "Generating draft for cluster: %s (%d firms, %d-day window)",
        saturation.subject_key,
        saturation.effective_firm_count,
        saturation.window_days,
    )

    # ── Step 1: fetch and enrich cluster articles ──────────────────────────────
    articles = _enrich_cluster_articles(saturation.cluster_id)
    if not articles:
        logger.error(
            "No articles found for cluster %d — cannot generate draft",
            saturation.cluster_id,
        )
        return None

    logger.info(
        "Enriched %d articles (%d with full text)",
        len(articles),
        sum(1 for a in articles if a.full_text),
    )

    # ── Step 2: build prompts ──────────────────────────────────────────────────
    system_prompt = _get_system_prompt()
    user_prompt   = _build_draft_prompt(saturation, articles)

    # ── Step 3: call Claude ────────────────────────────────────────────────────
    try:
        draft_text = _call_draft_api(system_prompt, user_prompt)
    except Exception as exc:
        logger.error(
            "Draft API call failed for %s: %s",
            saturation.subject_key, exc,
        )
        return None

    word_count = _count_words(draft_text)
    logger.info(
        "Draft generated: ~%d words for %s",
        word_count, saturation.subject_key,
    )

    # ── Step 4: build .docx ────────────────────────────────────────────────────
    try:
        docx_path = _parse_and_build_docx(saturation.subject_key, draft_text)
    except Exception as exc:
        logger.error(
            "Failed to build .docx for %s: %s",
            saturation.subject_key, exc,
        )
        return None

    return DraftResult(
        cluster_id=saturation.cluster_id,
        subject_key=saturation.subject_key,
        docx_path=docx_path,
        draft_text=draft_text,
        word_count=word_count,
    )
